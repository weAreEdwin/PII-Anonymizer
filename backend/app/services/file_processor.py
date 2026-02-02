import io
from typing import Tuple
from fastapi import UploadFile, HTTPException, status
import PyPDF2
from docx import Document


class FileProcessor:
    """Service for extracting text from various file formats"""
    
    ALLOWED_EXTENSIONS = {'.pdf', '.docx', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @staticmethod
    def validate_file(file: UploadFile) -> None:
        """Validate file type and size"""
        # Check file extension
        filename = file.filename.lower()
        if not any(filename.endswith(ext) for ext in FileProcessor.ALLOWED_EXTENSIONS):
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File type not allowed. Accepted types: {', '.join(FileProcessor.ALLOWED_EXTENSIONS)}"
            )
    
    @staticmethod
    async def extract_text(file: UploadFile) -> Tuple[str, str]:
        """
        Extract text from uploaded file
        
        Returns:
            Tuple of (extracted_text, file_type)
        """
        FileProcessor.validate_file(file)
        
        content = await file.read()
        
        # Check file size
        if len(content) > FileProcessor.MAX_FILE_SIZE:
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail=f"File size exceeds maximum allowed size of {FileProcessor.MAX_FILE_SIZE / (1024*1024)}MB"
            )
        
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.pdf'):
                text = FileProcessor._extract_from_pdf(content)
                return text, 'pdf'
            
            elif filename.endswith('.docx'):
                text = FileProcessor._extract_from_docx(content)
                return text, 'docx'
            
            elif filename.endswith('.txt'):
                text = content.decode('utf-8')
                return text, 'txt'
            
            else:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Unsupported file format"
                )
        
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Error processing file: {str(e)}"
            )
    
    @staticmethod
    def _extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file"""
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = '\n'.join(text_parts)
        
        if not full_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No text could be extracted from PDF. The file may be image-based or encrypted."
            )
        
        return full_text
    
    @staticmethod
    def _extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file"""
        docx_file = io.BytesIO(content)
        doc = Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = '\n'.join(text_parts)
        
        if not full_text.strip():
            raise HTTPException(
                status_code=status.HTTP_400_BAD_REQUEST,
                detail="No text could be extracted from DOCX file."
            )
        
        return full_text
